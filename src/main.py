import sys
import json
import random
import hashlib
import sqlite3
import winreg
import docx
import requests
import datetime
import re
import shutil
import http.client
import cv2
import os
import webbrowser
import globalvar
import configparser
import time
import urllib.request
import win32con, win32api
from urllib import parse
from pydub import AudioSegment
from pydub.playback import play
from aip import AipOcr
from PIL import ImageGrab
from ping3 import ping
from Widget import Ui_Dialog
from user import Ui_User
from register import Ui_reg
from findpasswd import Ui_fpasswd
from changeuser import Ui_cguser
from about import Ui_abt
from PyQt5.QtWidgets import QApplication, QWidget, QDialog, QTableWidgetItem, QAbstractItemView, \
    QHeaderView, QFileDialog, QMessageBox
from PyQt5.QtGui import QPixmap, QIcon
from PyQt5 import QtCore, Qt
from PyQt5.QtCore import QTimer


# 百度字符识别
class OCR:
    def __init__(self):
        self.AppID = config.get('ocr', 'appiD')
        self.API_Key = config.get('ocr', 'api_key')
        self.Secret_Key = config.get('ocr', 'secret_key')

    # 打开文件
    def getfilecontent(self, filePath):
        with open(filePath, 'rb') as f:
            return f.read()

    # 识别
    def identify(self, imgname, path):
        try:
            client = AipOcr(self.AppID, self.API_Key, self.Secret_Key)
            image = self.getfilecontent(path)
            res = client.basicGeneral(image)
            UI.textEdit.clear()
            for item in res['words_result']:
                UI.textEdit.append(item['words'])
            Widget().translate_()
        except:
            UI.textEdit.setText("截图识别失败: 请检查api有效性")
            errorbox.write(str(datetime.datetime.now()) + '\n')
            errorbox.write("截图识别失败: 请检查api有效性\n")
        # 删除临时图像
        if imgname == 'cut.jpg':
            os.remove('cut.jpg')


# 单词朗读
class pronounce:
    def __init__(self):
        self.aurl = 'http://dict.youdao.com/dictvoice?type=0&audio='
        self.uurl = 'http://dict.youdao.com/dictvoice?type=1&audio='

    # 拼接url
    def geturl(self, word):
        if config.get('pronounce', 'mode') == 'us':
            if os.path.exists('./voiceUS/' + word + '.mp3'):
                path = './voiceUS/' + word + '.mp3'
                self.playsound(path)
                return
            url = self.aurl + word
        elif config.get('pronounce', 'mode') == 'uk':
            if os.path.exists('./voiceUK/' + word + '.mp3'):
                path = './voiceUK/' + word + '.mp3'
                self.playsound(path)
                return
            url = self.uurl + word
        else:
            return
        self.download(url, word)

    # 从有道下载语音
    def download(self, url, word):
        if config.get('pronounce', 'mode') == 'us':
            filepath = './voiceUS/' + word.lower() + '.mp3'
        else:
            filepath = './voiceUK/' + word.lower() + '.mp3'

        filepath = filepath.replace("%20", " ")
        # 路径太长，重命名防止系统崩溃
        if len(filepath) > 200:
            filepath = 'tmp.mp3'
        urllib.request.urlretrieve(url, filepath)
        self.playsound(filepath)
        # 是否缓存
        if not config.getboolean('pronounce', 'cache'):
            os.remove(filepath)
        # 无论如何不保存长度超长的音频文件
        if filepath == 'tmp.mp3':
            os.remove(filepath)

    # 播放音频
    def playsound(self, file):
        try:
            # 标准mp3格式
            song = AudioSegment.from_mp3(file)
            play(song)
        except:
            # 非标准，先转换
            AudioSegment.from_file(file).export(file, format='mp3')
            song = AudioSegment.from_mp3(file)
            play(song)


class imgGrab(QWidget):

    # 抓图
    def cut(self):
        global img
        if not ping_163('www.163.com'):
            UI.textEdit.setText("网络故障: 请检查网络连通性")
            errorbox.write(str(datetime.datetime.now()) + '\n')
            errorbox.write("网络故障, 请检查网络连通性\n")
            return
        widget.showMinimized()
        # 窗口最小化等待0.3s，确保完成最小化
        time.sleep(0.3)
        self.scrren_cut()
        img = cv2.imread('screen.jpg')
        cv2.namedWindow('image')
        cv2.setMouseCallback('image', self.on_mouse)
        cv2.imshow('image', img)
        cv2.waitKey(0)
        os.remove('screen.jpg')

    # 抓取全屏
    def scrren_cut(self):
        image = ImageGrab.grab()
        image.save("screen.jpg")

    # 鼠标事件
    def on_mouse(self, event, x, y, flags, param):
        global img, point1, point2
        img2 = img.copy()
        if event == cv2.EVENT_LBUTTONDOWN:  # 左键点击
            point1 = (x, y)
            cv2.circle(img2, point1, 10, (255, 0, 0), 1)
            cv2.imshow('image', img2)
        elif event == cv2.EVENT_MOUSEMOVE and (flags & cv2.EVENT_FLAG_LBUTTON):  # 按住左键拖曳
            cv2.rectangle(img2, point1, (x, y), (255, 0, 0), 1)
            cv2.imshow('image', img2)
        elif event == cv2.EVENT_LBUTTONUP:  # 左键释放
            point2 = (x, y)
            cv2.rectangle(img2, point1, point2, (255, 0, 0), 1)
            cv2.imshow('image', img2)
            min_x = min(point1[0], point2[0])
            min_y = min(point1[1], point2[1])
            width = abs(point1[0] - point2[0])
            height = abs(point1[1] - point2[1])
            cut_img = img[min_y:min_y + height, min_x:min_x + width]
            cv2.imwrite('cut.jpg', cut_img)
            info = QMessageBox.warning(self, "截图信息", "截图完成？", QMessageBox.Ok | QMessageBox.Cancel)
            if info == QMessageBox.Ok:
                cv2.destroyWindow('image')
                widget.showNormal()
                ocr = OCR()
                ocr.identify('cut.jpg', './cut.jpg')


class BaiduTranslate:
    def __init__(self, fromLang, toLang):
        self.url = "/api/trans/vip/translate"
        self.appid = config.get('translate', 'appid')
        self.secretKey = config.get('translate', 'secretkey')
        self.fromLang = fromLang
        self.toLang = toLang
        self.salt = random.randint(32768, 65536)

    # 百度翻译demo，向百度请求并解析返回值
    def BdTrans(self, text):
        sign = self.appid + text + str(self.salt) + self.secretKey
        md = hashlib.md5()
        md.update(sign.encode(encoding='utf-8'))
        sign = md.hexdigest()
        myurl = self.url + \
                '?appid=' + self.appid + \
                '&q=' + parse.quote(text) + \
                '&from=' + self.fromLang + \
                '&to=' + self.toLang + \
                '&salt=' + str(self.salt) + \
                '&sign=' + sign
        try:
            httpClient = http.client.HTTPConnection('api.fanyi.baidu.com')
            httpClient.request('GET', myurl)
            response = httpClient.getresponse()
            html = response.read().decode('utf-8')
            html = json.loads(html)
            dst = html["trans_result"][0]["dst"]
            return dst
        except Exception as e:
            errorbox.write(str(datetime.datetime.now()) + '\n')
            if not ping_163('www.163.com'):
                errorbox.write("网络故障: 请检查网络连通性")
                return "网络故障: 请检查网络连通性"
            else:
                errorbox.write("翻译失败: 请检查api有效性")
                return "翻译失败: 请检查api有效性"


class changeuser(QDialog, Ui_cguser):
    def __init__(self):
        super(changeuser, self).__init__()
        self.ui = Ui_cguser()
        self.ui.setupUi(self)
        # 隐藏标题栏
        self.setWindowFlags(Qt.Qt.FramelessWindowHint)
        self.ui.pushButton.clicked.connect(self.check)
        self.ui.pushButton.setStyleSheet('QPushButton:hover{background:rgb(31,199,253);'
                                         'border-radius:4px;}'
                                         'QPushButton{color:#ffffff}'
                                         'QPushButton{background:rgb(7,189,253)}'
                                         'QPushButton{border-radius:4px;}')
        self.ui.pushButton_2.setStyleSheet('QPushButton:hover{background:red;'
                                           'border-radius:2px;}'
                                           'QPushButton{background:transparent;}')
        self.ui.pushButton_2.setDefault(False)
        self.ui.pushButton_3.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                           'border-radius:2px;}'
                                           'QPushButton{background:transparent;}')
        self.ui.pushButton_3.setDefault(False)

    # 切换用户，检测密码密码是否正确
    def check(self):
        sql = "SELECT password FROM user WHERE username=?"
        c.execute(sql, (UI.comboBox.currentText(),))
        password = c.fetchall()[0][0]
        if password == self.ui.lineEdit.text():
            QMessageBox.information(self, "切换用户", "已切换用户为{}".format(UI.comboBox.currentText()))
            # 修改当前用户索引
            config.set('user', 'cur', str(UI.comboBox.currentIndex()))
            config.write(open('config.ini', 'w'))
            self.accept()
        else:
            QMessageBox.warning(self, "切换用户", "密码错误")


class findpasswd(QDialog, Ui_fpasswd):
    def __init__(self):
        super(findpasswd, self).__init__()
        self.ui = Ui_fpasswd()
        self.ui.setupUi(self)
        # 隐藏标题栏
        self.setWindowFlags(Qt.Qt.FramelessWindowHint)
        self.ui.pushButton.clicked.connect(self.gettext)
        self.ui.pushButton_2.setStyleSheet('QPushButton:hover{background:red;'
                                           'border-radius:2px;}'
                                           'QPushButton{background:transparent;}')
        self.ui.pushButton_3.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                           'border-radius:2px;}'
                                           'QPushButton{background:transparent;}')
        for index in c.execute("SELECT * FROM user"):
            self.ui.comboBox.addItem(index[0])
        # 设置相应用户的密保问题
        self.ui.comboBox.currentTextChanged.connect(self.setlineedit)
        # 禁用密保问题输入框，由程序查数据库输入
        self.ui.lineEdit.setEnabled(False)

    # 获取用户输入，验证两次密码输入是否相同
    def gettext(self):
        username = self.ui.comboBox.currentText()
        answer = self.ui.lineEdit_2.text()
        newpswd = self.ui.lineEdit_3.text()
        cpswd = self.ui.lineEdit_4.text()
        if newpswd != cpswd:
            QMessageBox.warning(self, "找回密码", "两次输入的密码不一致")
            return
        self.check(username, answer, newpswd)

    # 找回密码验证合法性
    def check(self, username, answer, newpswd):
        sql = 'SELECT * FROM user where username = ?'
        c.execute(sql, (username,))
        if answer == c.fetchall()[0][3]:
            p = re.compile(r'^(?![0-9]+$)(?![a-zA-Z]+$)[0-9A-Za-z]{6,12}$')
            if p.match(newpswd):
                sql = 'UPDATE user SET password=? where username=?'
                c.execute(sql, (newpswd, username,))
                conn.commit()
                QMessageBox.information(self, "找回密码", "密码已重设")
                self.accept()
            else:
                QMessageBox.warning(self, "找回密码", "新密码不符合要求")
        else:
            QMessageBox.warning(self, "找回密码", "密保问题答案错误")

    # 设置显示密保问题
    def setlineedit(self):
        sql = 'SELECT * FROM user where username = ?'
        c.execute(sql, (self.ui.comboBox.currentText(),))
        self.ui.lineEdit.setText(c.fetchall()[0][2])


class register(QDialog, Ui_reg):
    def __init__(self):
        super(register, self).__init__()
        self.ui = Ui_reg()
        self.ui.setupUi(self)
        self.setWindowFlags(Qt.Qt.FramelessWindowHint)
        self.setbtns()
        self.showabout()
        self.picname = ""
        self.timer = QTimer()
        self.timer.start(5000)
        self.timer.timeout.connect(self.changbg)
        self.counter = 0

    # 自动显示注册注意事项
    def showabout(self):
        if config.getboolean('about', 'show'):
            about = About()
            about.exec()

    # 设置按钮样式表，连接信号和槽
    def setbtns(self):
        self.ui.pushButton.clicked.connect(self.getimg)
        self.ui.pushButton_2.clicked.connect(self.gettext)
        self.ui.pushButton_5.clicked.connect(self.showabout_)
        self.ui.pushButton_3.setStyleSheet('QPushButton:hover{background:red;'
                                           'border-radius:2px;}'
                                           'QPushButton{background:transparent;}')
        self.ui.pushButton_4.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                           'border-radius:2px;}'
                                           'QPushButton{background:transparent;}')
        self.ui.pushButton_5.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                           'border-radius:2px;}'
                                           'QPushButton{background:transparent;}')
        self.ui.label.setStyleSheet("image: url(:/pictures/img/regbg.png);")

    # 点击显示注册注意事项
    def showabout_(self):
        about = About()
        about.exec()

    # 打开图片
    def getimg(self):
        imgName, imgType = QFileDialog.getOpenFileName(self, "打开图片", "", "*.jpg;*.png;;All Files(*)")
        self.picname = imgName
        jpg = QPixmap(self.picname).scaled(self.ui.label_5.width(), self.ui.label_5.height())
        self.ui.label_5.setStyleSheet("")
        self.ui.label_5.setPixmap(jpg)

    # 每5秒更新背景，5张图片循环
    def changbg(self):
        if self.counter == 0:
            self.ui.label.setStyleSheet("image: url(:/pictures/img/regbg3.png);")
            self.ui.label.repaint()
            self.counter = 1
        elif self.counter == 1:
            self.ui.label.setStyleSheet("image: url(:/pictures/img/regbg2.png);")
            self.ui.label.repaint()
            self.counter = 2
        elif self.counter == 2:
            self.ui.label.setStyleSheet("image: url(:/pictures/img/regbg4.png);")
            self.ui.label.repaint()
            self.counter = 3
        elif self.counter == 3:
            self.ui.label.setStyleSheet("image: url(:/pictures/img/regbg.png);")
            self.ui.label.repaint()
            self.counter = 4
        elif self.counter == 4:
            self.ui.label.setStyleSheet("image: url(:/pictures/img/regbg5.png);")
            self.ui.label.repaint()
            self.counter = 0

    # 获取用户输入
    def gettext(self):
        username = self.ui.lineEdit.text()
        password = self.ui.lineEdit_2.text()
        password_ = self.ui.lineEdit_5.text()
        question = self.ui.lineEdit_3.text()
        answer = self.ui.lineEdit_4.text()
        self.check(username, password, password_, question, answer)

    # 用户注册检测合法性
    def check(self, username, password, password_, question, answer):
        if self.picname == "":
            QMessageBox.warning(None, "用户注册", "请点击上方图标以选择用户头像")
            self.getimg()
            return
        if username == "" or password == "":
            QMessageBox.warning(None, "用户注册", "用户名和密码不能为空")
            return
        if password != password_:
            QMessageBox.warning(None, "用户注册", "两次输入密码不一致")
            return
        else:
            for index in c.execute("SELECT * FROM user"):
                if index[0] == username:
                    QMessageBox.warning(None, "用户注册", "该用户名已存在")
                    return
        u = re.compile(r'^[A-Z][\w]{1,8}$')
        # 由大写字母开头，只能有字母数字，1-8位
        p = re.compile(r'^(?![0-9]+$)(?![a-zA-Z]+$)[0-9A-Za-z]{6,12}$')
        # 由数字和字母组成，并且要同时含有数字和字母，且长度要在6-12位之间。
        if p.match(password) and u.match(username):
            try:
                path, name = os.path.splitext(self.picname)
                temp = self.picname
                self.picname = "./userportrait/" + os.path.join(username) + os.path.join(name)
                shutil.copyfile(temp, self.picname)
            except:
                errorbox.write(str(datetime.datetime.now()) + '\n')
                errorbox.write("目标图像已存在")
            # 新用户信息插入数据库
            c.execute("INSERT INTO user VALUES(?,?,?,?,?)", (username, password, question, answer, self.picname))
            conn.commit()
            globalvar.CUR_USER = username
            sql = 'CREATE TABLE {}(src TEXT, dst TEXT)'.format(globalvar.CUR_USER)
            c.execute(sql)
            conn.commit()
            # 防止程序崩溃，暂时取消信号和槽连接
            UI_U.comboBox.currentTextChanged.disconnect(user.setportrait)
            UI_U.comboBox.clear()
            # 重新设置combobox选单
            for index in c.execute('SELECT username FROM user'):
                UI_U.comboBox.addItem(index[0])
            UI_U.comboBox.currentTextChanged.connect(user.setportrait)
            UI_U.comboBox.setCurrentText(username)
            # UI_U.lineEdit_2.setText(password)
            self.accept()
        else:
            QMessageBox.warning(None, "用户注册", "用户名由大写字母开头，可包含字母和数字，长度限制1-8位,\n密码需同时且只能包含字母和数字，长度限制6-12位")


class About(QDialog, Ui_abt):
    def __init__(self):
        super(About, self).__init__()
        self.ui = Ui_abt()
        self.ui.setupUi(self)
        if not config.getboolean('about', 'show'):
            self.ui.checkBox.setChecked(True)
        self.ui.checkBox.stateChanged.connect(self.setstate)

    # 设置不再显示
    def setstate(self):
        if self.ui.checkBox.isChecked():
            config.set('about', 'show', 'False')
        else:
            config.set('about', 'show', 'True')

        config.write(open('config.ini', 'w'))


class User(QDialog, Ui_User):
    def __init__(self):
        super(User, self).__init__()
        self.ui = Ui_User()
        self.ui.setupUi(self)
        self.setWindowFlags(Qt.Qt.FramelessWindowHint)
        # 隐藏两个frame
        self.ui.frame.setVisible(False)
        self.ui.frame_2.setVisible(False)
        self.tmp_check = 0
        self.tmp_check_ = 0
        self.sec = 0
        self.conbocount = -1
        self.timer = QTimer()
        try:
            self.timeout = config.getint('user', 'timeout')
            globalvar.rempasswd = config.getboolean('config', 'rempasswd')
            globalvar.autologin = config.getboolean('config', 'autologin')
        except:
            QMessageBox.warning(None, "错误", "找不到配置文件,程序将退出")
            exit(1)
        if globalvar.rempasswd:
            self.ui.checkBox.setChecked(True)
        if globalvar.autologin and config.getboolean('user', 'exist'):
            self.ui.checkBox_2.setChecked(True)
            self.ui.checkBox.setChecked(True)
            self.ui.pushButton.setText(str(self.timeout) + "秒后自动登录")

    # 自动登录，定时器超时函数
    def autologin(self):
        if self.sec == self.timeout:
            self.getinfo()
        self.sec += 1
        text = str(self.timeout - self.sec) + "秒后自动登录"
        self.ui.pushButton.setText(text)

    # user类初始化
    def init(self):
        self.setbtns()
        for index in c.execute("SELECT * FROM user"):
            UI_U.comboBox.addItem(index[0])
            self.conbocount += 1
        # 自动登录并且用户索引合法
        if config.getint('user', 'cur') <= self.conbocount and config.getboolean('config', 'autologin'):
            self.timer.start(1000)
            self.timer.timeout.connect(self.autologin)
        UI_U.comboBox.currentTextChanged.connect(self.setportrait)
        UI_U.checkBox.stateChanged.connect(self.setstate)
        UI_U.checkBox_2.stateChanged.connect(self.setstate2)
        if config.getint('user', 'cur') <= self.conbocount:
            UI_U.comboBox.setCurrentIndex(config.getint('user', 'cur'))
        if config.getboolean('user', 'exist'):
            self.setportrait()

    # 设置按钮样式表，连接信号和槽，设置config
    def setbtns(self):
        UI_U.pushButton.clicked.connect(self.getinfo)
        UI_U.pushButton.setStyleSheet('QPushButton:hover{background:rgb(31,199,253);'
                                      'border-radius:4px;}'
                                      'QPushButton{color:#ffffff}'
                                      'QPushButton{background:rgb(7,189,253)}'
                                      'QPushButton{border-radius:4px;}')
        UI_U.pushButton.setDefault(False)
        UI_U.pushButton_2.setStyleSheet('QPushButton:hover{background:red;'
                                        'border-radius:2px;}'
                                        'QPushButton{background:transparent;}')

        UI_U.pushButton_3.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                        'border-radius:2px;}'
                                        'QPushButton{background:transparent;}')
        UI_U.pushButton_3.setDefault(False)

        UI_U.pushButton_4.setStyleSheet('QPushButton:hover{color:rgb(131,131,131);}'
                                        'QPushButton{color:#b2b2b2}'
                                        'QPushButton{background:transparent;}')
        UI_U.pushButton_4.clicked.connect(User.register)
        UI_U.pushButton_6.setStyleSheet('QPushButton:hover{color:rgb(131,131,131);}'
                                        'QPushButton{color:#b2b2b2}'
                                        'QPushButton{background:transparent;}')
        UI_U.pushButton_7.setStyleSheet('QPushButton:hover{background:rgb(31,199,253);'
                                        'border-radius:4px;}'
                                        'QPushButton{color:#ffffff}'
                                        'QPushButton{background:rgb(7,189,253)}'
                                        'QPushButton{border-radius:4px;}')
        UI_U.pushButton_8.setStyleSheet('QPushButton:hover{background:red;'
                                        'border-radius:2px;}'
                                        'QPushButton{background:transparent;}')
        UI_U.pushButton_9.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                        'border-radius:2px;}'
                                        'QPushButton{background:transparent;}')
        UI_U.pushButton_10.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                         'border-radius:3px;}'
                                         'QPushButton{background:transparent;}')
        UI_U.pushButton_11.setStyleSheet('QPushButton{background:rgb(244,244,244);'
                                         'border-radius:2px;'
                                         'border:3px;}'
                                         'QPushButton:hover{background:rgb(190,231,253);'
                                         'border-radius:4px;}')
        UI_U.pushButton_12.setStyleSheet('QPushButton{background:rgb(244,244,244);'
                                         'border-radius:2px;'
                                         'border:1px;}'
                                         'QPushButton:hover{background:rgb(190,231,253);'
                                         'border-radius:4px;}')
        UI_U.pushButton_13.setStyleSheet('QPushButton:hover{background:red;'
                                         'border-radius:2px;}'
                                         'QPushButton{background:transparent;}')
        UI_U.pushButton_14.setStyleSheet('QPushButton:hover{background:rgb(220,220,220);'
                                         'border-radius:2px;}'
                                         'QPushButton{background:transparent;}')
        UI_U.pushButton_6.clicked.connect(self.resetpasswd)
        UI_U.pushButton_5.clicked.connect(self.setframeshow)
        UI_U.pushButton_7.clicked.connect(self.setframehide)
        UI_U.pushButton_10.clicked.connect(self.setconfigshow)
        UI_U.pushButton_11.clicked.connect(self.setconfighide)
        UI_U.pushButton_12.clicked.connect(self.setconfig)
        # 从config.ini获取值
        UI_U.lineEdit.setText(str(config.get('ocr', 'appid')))
        UI_U.lineEdit_3.setText(str(config.get('ocr', 'api_key')))
        UI_U.lineEdit_4.setText(str(config.get('ocr', 'secret_key')))
        UI_U.lineEdit_5.setText(str(config.get('translate', 'appid')))
        UI_U.lineEdit_6.setText(str(config.get('translate', 'secretkey')))
        UI_U.lineEdit_7.setText(str(config.get('user', 'timeout')))
        UI_U.lineEdit_8.setText(str(config.get('pronounce', 'mode')))
        UI_U.lineEdit_9.setText(str(config.get('pronounce', 'cache')))

    # 从用户输入设置config文件
    def setconfig(self):
        config.set('ocr', 'appid', UI_U.lineEdit.text())
        config.set('ocr', 'api_key', UI_U.lineEdit_3.text())
        config.set('ocr', 'secret_key', UI_U.lineEdit_4.text())
        config.set('translate', 'appid', UI_U.lineEdit_5.text())
        config.set('translate', 'secretkey', UI_U.lineEdit_6.text())

        # 确保设置倒计时符合规则，防止程序崩溃
        if re.compile(r'^[0-9]{1,3}$').match(UI_U.lineEdit_7.text()):
            config.set('user', 'timeout', UI_U.lineEdit_7.text())
        else:
            config.set('user', 'timeout', '3')
        # 确保音种设置符合规则，防止程序崩溃
        if UI_U.lineEdit_8.text().lower() == 'us' or UI_U.lineEdit_8.text().lower() == 'uk':
            config.set('pronounce', 'mode', UI_U.lineEdit_8.text())
        else:
            config.set('pronounce', 'mode', 'uk')
        # 确保缓存设置符合规则
        if UI_U.lineEdit_9.text().lower() == 'true' or UI_U.lineEdit_9.text().lower() == 'false':
            config.set('pronounce', 'cache', UI_U.lineEdit_9.text())
        else:
            config.set('pronounce', 'cache', 'true')
        config.write(open('config.ini', 'w'))
        self.setconfighide()

    # 设置frame显示和隐藏
    def setconfigshow(self):
        UI_U.frame_2.setVisible(True)

    def setconfighide(self):
        UI_U.frame_2.setVisible(False)

    def setframeshow(self):
        UI_U.frame.setVisible(True)

    def setframehide(self):
        UI_U.frame.setVisible(False)

    # 勾选记住密码
    def setstate(self):
        if UI_U.checkBox.isChecked():
            # 暂时记住勾选，登录成功后确认勾选
            self.tmp_check = 1
            # config.set('config', 'rempasswd', 'True')
        else:
            # 取消勾选
            self.tmp_check = 0
            config.set('config', 'rempasswd', 'False')
            UI_U.checkBox_2.setChecked(False)
        config.write(open('config.ini', 'w'))

    # 勾选自动登录
    def setstate2(self):
        if UI_U.checkBox_2.isChecked():
            # 暂时记住勾选，登录成功后确认勾选
            self.tmp_check_ = 1
            # config.set('config', 'rempasswd', 'True')
            # config.set('config', 'autologin', 'True')
            UI_U.checkBox.setChecked(True)
        else:
            # 取消勾选
            self.tmp_check_ = 0
            self.timer.stop()
            self.ui.pushButton.setText("登录")
            config.set('config', 'autologin', 'False')
        config.write(open('config.ini', 'w'))

    # 用户注册
    def register(self):
        reg = register()
        # 如果自动登录，取消倒计时
        user.timer.stop()
        user.ui.pushButton.setText("登录")
        if reg.exec():
            if not config.getboolean('user', 'exist'):
                user.setportrait()

    # 忘记密码
    def resetpasswd(self):
        fp = findpasswd()
        # 如果自动登录，取消倒计时
        user.timer.stop()
        user.ui.pushButton.setText("登录")
        fp.exec()

    # 设置头像
    def setportrait(self):
        sql = 'SELECT portrait FROM user where username=?'
        if not config.getboolean('user', 'exist'):
            c.execute(sql, (globalvar.CUR_USER,))
        else:
            c.execute(sql, (self.ui.comboBox.currentText(),))
        path = c.fetchall()[0][0]
        jpg = QPixmap(path).scaled(self.ui.label_5.width(), self.ui.label_5.height())
        self.ui.label_5.setStyleSheet("background:transparent")
        self.ui.label_5.setPixmap(jpg)
        self.ui.label_5.repaint()
        # 如果记住密码，则自动填写密码
        if globalvar.rempasswd:
            sql = 'SELECT password FROM user where username=?'
            c.execute(sql, (self.ui.comboBox.currentText(),))
            password = c.fetchall()[0][0]
            self.ui.lineEdit_2.setText(password)
        else:
            self.ui.lineEdit_2.clear()

    # 从文本框获取用户名密码
    def getinfo(self):
        username = UI_U.comboBox.currentText()
        password = UI_U.lineEdit_2.text()
        if username != "" and password != "":
            self.check(username, password)
        else:
            UI_U.checkBox.setChecked(False)
            UI_U.checkBox_2.setChecked(False)
            QMessageBox.warning(None, "出错啦", "未输入用户名或密码")

    # 检测用户名和密码是否正确
    def check(self, username, password):
        for index in c.execute("SELECT * FROM user"):
            # 验证存在用户名
            if index[0] == username:
                # 密码正确
                if index[1] == password:
                    globalvar.CUR_USER = username
                    config.set('user', 'cur', str(self.ui.comboBox.currentIndex()))
                    if self.tmp_check == 1:
                        config.set('config', 'rempasswd', 'True')
                    if self.tmp_check_ == 1:
                        config.set('config', 'autologin', 'True')
                    config.write(open('config.ini', 'w'))
                    self.accept()
                else:
                    QMessageBox.warning(None, "warning", "密码错误")
                    break


class Widget(QWidget, Ui_Dialog):
    def __init__(self):
        super(Widget, self).__init__()
        self.apptimer = QTimer()
        self.ui = Ui_Dialog()
        self.ui.setupUi(self)
        self.setFixedSize(self.width(), self.height())
        self.counter = 0
        self.timecounter = 0
        self.todaytimecounter = config.getint('statistics', 'today')
        self.totaltimecounter = config.getint('statistics', 'total')
        self.date = config.getint('statistics', 'date')

    # 统计
    def statistics(self):
        # check date
        if self.date != int(time.time()/86400):
            self.date = int(time.time()/86400)
            self.todaytimecounter = 0

        self.timecounter += 1
        # this time
        m, s = divmod(self.timecounter, 60)
        h, m = divmod(m, 60)
        str = ("%02d:%02d:%02d" % (h, m, s))
        self.ui.lcdNumber_4.display(str)

        # today
        m, s = divmod(self.timecounter+self.todaytimecounter, 60)
        h, m = divmod(m, 60)
        str = ("%02d:%02d:%02d" % (h, m, s))
        self.ui.lcdNumber_5.display(str)

        # total
        m, s = divmod(self.totaltimecounter+self.timecounter, 60)
        h, m = divmod(m, 60)
        str = ("%02d:%02d:%02d" % (h, m, s))
        self.ui.lcdNumber_2.display(str)


    # 初始化
    def init(self):
        self.apptimer.start(1000)
        self.apptimer.timeout.connect(self.statistics)
        cur = datetime.datetime.now()
        if cur.day<=9:
            UI.label_3.setText('0'+str(cur.day))
        else:
            UI.label_3.setText('0'+str(cur.day))
        # 月份数字转英文缩写
        UI.label_4.setText(
            {1: 'Jan.',
             2: 'Feb.',
             3: 'Mar.',
             4: 'Apr.',
             5: 'May.',
             6: 'Jun.',
             7: 'Jul.',
             8: 'Aug.',
             9: 'Sep.',
             10: 'Oct.',
             11: 'Nov.',
             12: 'Dec.'}.get(cur.month))

        for username in c.execute("SELECT * FROM user"):
            UI.comboBox.addItem(QIcon(username[4]), username[0])

        for index in c.execute("SELECT * FROM user"):
            globalvar.counter += 1
            if globalvar.CUR_USER == index[0]:
                break
        UI.comboBox.setCurrentIndex(globalvar.counter - 1)
        # 输入框clear按钮
        UI.lineEdit.setClearButtonEnabled(True)
        UI.tableWidget.setEditTriggers(QAbstractItemView.NoEditTriggers)
        UI.tableWidget.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        widget.setWindowFlags(QtCore.Qt.WindowMinimizeButtonHint)
        self.ui.tabWidget.setCurrentIndex(0)
        self.signalslot()

    # 源语言
    def srclang(self):
        src = UI.buttonGroup.checkedButton().text()
        if src == '自动检测':
            return 'auto'
        else:
            return 'zh'

    # 目标语言
    def dstlang(self):
        dst = UI.buttonGroup.checkedButton().text()
        text = UI.lineEdit.text()
        if dst == '自动检测' and 'a' <= text.lower()[0] <= 'z':
            return 'zh'
        else:
            return {
                '汉译英': 'en',
                '汉译法': 'fra',
                '汉译日': 'jp',
                '外译汉': 'zh',
                '汉译俄': 'ru',
                '汉译韩': 'kor',
                '汉译西': 'spa',
                '汉译德': 'de',
                '汉译粤': 'yue',
                '简译繁': 'cht',
                '汉译古': 'wyw',
                '汉译意': 'it',
                '自动检测': 'en'
            }.get(dst)

    # 清空数据库和收藏夹
    def clear(self):
        sql = 'DELETE FROM {}'.format(globalvar.CUR_USER)
        c.execute(sql)
        conn.commit()
        UI.tableWidget.setRowCount(0)
        UI.tableWidget.setHorizontalHeaderLabels(['源语言', '目标语言'])
        UI.tableWidget.setColumnCount(2)

    # 从数据库中删除，更新收藏夹
    def delete(self):
        num = UI.tableWidget.currentRow()
        if num <= UI.tableWidget.rowCount():
            try:
                p = UI.tableWidget.item(num, 0).text()
                sql = "DELETE FROM " + globalvar.CUR_USER + " where src=" + "'" + p + "'"
                c.execute(sql)
                conn.commit()
                self.updateTable()
            except:
                QMessageBox.warning(self,"waring", "请选择项")

    # 选择文档翻译
    def openfile(self):
        file, filetype = QFileDialog.getOpenFileName(self, "open", self.desktop_path(), "*.docx")
        try:
            file = docx.Document(file)
            UI.textEdit.clear()
            for contents in file.paragraphs:
                UI.textEdit.insertPlainText(contents.text)
            self.translate_()
            UI.textEdit.clear()
            for contents in file.paragraphs:
                UI.textEdit.append(contents.text)

        except:
            print("open canceled")

    # 人工翻译
    def openyoudao(self):
        url = 'https://f.youdao.com/?path=fast&keyfrom=Nav-fast'
        webbrowser.open_new_tab(url)

    # 打开四六级官网
    def open46(self):
        url = 'http://cet.neea.edu.cn/'
        webbrowser.open_new_tab(url)

    # 切换用户
    def changeuser(self):
        UI.comboBox.currentTextChanged.disconnect(widget.changeuser)
        cguser = changeuser()
        if cguser.exec():
            globalvar.CUR_USER = self.ui.comboBox.currentText()
            widget.updateTable()
        else:
            UI.comboBox.setCurrentIndex(globalvar.counter - 1)
            UI.comboBox.currentTextChanged.connect(widget.changeuser)

    # 播放单词朗读语音
    def play(self):
        # 确保网络连通
        if ping_163('www.163.com'):
            prnc = pronounce()
            text = UI.textBrowser.toPlainText()
            text = re.sub(r'[^a-zA-Z]', "%20", text)
            if 'a' <= text[0].lower() <= 'z':
                prnc.geturl(text)
            else:
                QMessageBox.information(self, "词汇朗读", "抱歉，目前只支持英文朗读")
        else:
            UI.textBrowser.append("网络故障，单词朗读失败")
            errorbox.write(str(datetime.datetime.now()) + '\n')
            errorbox.write("网络故障，单词朗读失败")

    # 连接信号和槽
    def signalslot(self):
        UI.pushButton.clicked.connect(self.translate)
        UI.pushButton_3.clicked.connect(self.close_)
        UI.pushButton_2.clicked.connect(self.insertToDB)
        UI.pushButton_4.clicked.connect(self.openfile)
        UI.pushButton_5.clicked.connect(self.clear)
        UI.pushButton_6.clicked.connect(self.delete)
        UI.pushButton_7.clicked.connect(sc.cut)
        UI.pushButton_8.clicked.connect(self.translate_)
        UI.pushButton_9.clicked.connect(self.openyoudao)
        UI.pushButton_10.clicked.connect(self.open46)
        UI.pushButton_11.clicked.connect(self.imgidentify)
        UI.pushButton_12.clicked.connect(self.play)
        UI.pushButton_13.clicked.connect(loadsentences)
        UI.comboBox.currentTextChanged.connect(widget.changeuser)

    # 桌面路径
    def desktop_path(self):
        key = winreg.OpenKey(winreg.HKEY_CURRENT_USER,
                             r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
        path = winreg.QueryValueEx(key, "Desktop")[0]
        return path

    # 选择图片
    def imgidentify(self):
        try:
            imgName, imgType = QFileDialog.getOpenFileName(self, "打开图片", self.desktop_path(),
                                                           "*.jpg;*.png;;All Files(*)")
            path, name = os.path.split(imgName)
            ocr = OCR()
            ocr.identify(name, imgName)
        except:
            print("open canceled")

    # 更新收藏夹
    def updateTable(self):
        UI.tableWidget.setRowCount(0)
        UI.tableWidget.setHorizontalHeaderLabels(['源语言', '目标语言'])
        UI.tableWidget.setColumnCount(2)
        i = 0
        sql = 'SELECT * FROM {0}'.format(globalvar.CUR_USER)
        for str in c.execute(sql):
            if str == "":
                break
            UI.tableWidget.setRowCount(i + 1)
            UI.tableWidget.setItem(i, 0, QTableWidgetItem(str[0]))
            UI.tableWidget.setItem(i, 1, QTableWidgetItem(str[1]))
            i = i + 1

    # 单词翻译
    def translate(self):
        contents = UI.lineEdit.text()
        # 翻译内容不为空
        if not contents == "":
            # 获取源语言和目标语言
            BaiduTrans = BaiduTranslate(self.srclang(), self.dstlang())
            # 请求百度并解析返回值
            Results = BaiduTrans.BdTrans(contents)
            UI.textBrowser.setText(Results)

    # 检测下拉列表语种
    def langdetect(self):
        return {
            '中英互译': 'en',
            '中法互译': 'fra',
            '中德互译': 'de',
            '中日互译': 'jp',
            '中俄互译': 'ru'
        }.get(UI.comboBox_2.currentText())

    # 长文翻译
    def translate_(self):
        contents = UI.textEdit.toPlainText()
        # 要翻译的内容不为空
        if not contents == "":
            # 源语言是中文
            if re.compile(u'[\u4e00-\u9fa5]+').match(contents) is not None:
                BaiduTrans = BaiduTranslate('zh', self.langdetect())
            else:
                BaiduTrans = BaiduTranslate(self.langdetect(), 'zh')
            Results = BaiduTrans.BdTrans(contents)
            UI.textBrowser_2.setText(Results)

    # 数据库备份
    def dbbkp(self):
        shutil.copyfile('information.db', './backup/information.dbbkp')
        '''win32api.SetFileAttributes('./backup/information.dbbkp', win32con.FILE_ATTRIBUTE_HIDDEN
                                   | win32con.FILE_ATTRIBUTE_READONLY)'''

    # 关闭程序
    def close_(self):
        info = QMessageBox.warning(self, "Warning", "确定退出吗",
                                   QMessageBox.Ok | QMessageBox.Cancel)
        if info == QMessageBox.Ok:
            append = datetime.datetime.now()

            config.set('statistics', 'this', str((append - appstart).total_seconds()))
            config.set('statistics', 'today', str(self.todaytimecounter+self.timecounter))
            config.set('statistics', 'total', str(self.totaltimecounter+self.timecounter))
            config.set('statistics', 'date', str(int(time.time()/86400)))
            config.write(open('config.ini', 'w'))
            self.dbbkp()
            conn.close()
            errorbox.close()
            widget.close()

    # 插入单词到数据库
    def insertToDB(self):
        # 确保翻译成功
        if not UI.lineEdit.text() == "" and (not UI.textBrowser.toPlainText() == '翻译失败: 请检查api有效性' and
                                             not UI.textBrowser.toPlainText() == '网络故障: 请检查网络连通性'):
            sql = "INSERT INTO" + " " + globalvar.CUR_USER + " " + "VALUES" + " " + "('" + UI.lineEdit.text() + "'" + "," + "'" + UI.textBrowser.toPlainText() + "')"
            c.execute(sql)
            conn.commit()
            # 更新收藏夹
            self.updateTable()


# 检测网络连通性
def ping_163(ip):
    resp = ping(ip)
    if resp is not None:
        return True
    else:
        return False


# 加载每日一句，设置刷新按钮
def loadsentences():
    url = "http://open.iciba.com/dsapi/"
    try:
        r = requests.get(url)
        contents = r.json()['content']
        note = r.json()['note']
        UI.label.setText(contents)
        UI.label_2.setText(note)
        UI.label_2.setGeometry(20, 400, 401, 21)
        UI.pushButton_13.setVisible(False)

    except:
        UI.label.setText("Network error")
        UI.label_2.setText("网络故障，请检查网络连通性或点击左侧图标刷新")
        UI.pushButton_13.setGeometry(20, 400, 23, 23)
        UI.label_2.setGeometry(43, 400, 401, 21)
        UI.pushButton_13.setVisible(True)


# 隐藏文件，不对用户可见
def hidefiles():
    win32api.SetFileAttributes('information.db', win32con.FILE_ATTRIBUTE_HIDDEN)
    win32api.SetFileAttributes('errorbox', win32con.FILE_ATTRIBUTE_HIDDEN)
    win32api.SetFileAttributes('backup', win32con.FILE_ATTRIBUTE_HIDDEN)
    win32api.SetFileAttributes('voiceUS', win32con.FILE_ATTRIBUTE_HIDDEN)
    win32api.SetFileAttributes('voiceUK', win32con.FILE_ATTRIBUTE_HIDDEN)
    win32api.SetFileAttributes('userportrait', win32con.FILE_ATTRIBUTE_HIDDEN)


# 建立文件和文件夹
def mkfiles():
    if not os.path.exists('./voiceUS'):
        os.mkdir('./voiceUS')
    if not os.path.exists('./voiceUK'):
        os.mkdir('./voiceUK')
    if not os.path.exists('./backup'):
        os.mkdir('./backup')
    if not os.path.exists('./config.ini'):
        config['waring'] = {'waring': 'Do not try to modify config, or you understand what it means'
                            }
        config['ocr'] = {'appid': '',
                         'api_key': '',
                         'secret_key': ''
                         }
        config['translate'] = {'appid': '',
                               'secretkey': ''

                               }
        config['config'] = {'rempasswd': 'False',
                            'autologin': 'False'
                            }
        config['user'] = {'cur': '0',
                          'timeout': '5',
                          'exist': 'False'
                          }
        config['about'] = {'show': 'True'
                           }
        config['pronounce'] = {'mode': 'us',
                               'cache': 'True'
                               }
        config['statistics'] = {'total': '0',
                                'today': '0',
                                'this': '0',
                                'date': '0'}
        with open('config.ini', 'w') as configfile:
            config.write(configfile)
    if not os.path.exists('./userportrait'):
        os.mkdir('./userportrait')
    if not os.path.exists('./errorbox'):
        fp = open('./errorbox', 'w')
        fp.write("file created\n")
        fp.close()
    # 隐藏文件
    hidefiles()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    # 连接数据库
    conn = sqlite3.connect('information.db')
    c = conn.cursor()
    config = configparser.ConfigParser()
    mkfiles()
    config.read('config.ini')
    # 记录出错信息
    errorbox = open('./errorbox', 'a')
    try:
        c.execute("CREATE TABLE user(username text, password text, question text, answer text, portrait text)")
        conn.commit()
        # 用户第一次使用，注册时自动弹窗
        config.set('about', 'show', 'True')
        config.set('user', 'exist', 'False')
        config.write(open('config.ini', 'w'))

    except Exception as e:
        errorbox.write(str(datetime.datetime.now()) + '\n')
        errorbox.write("数据库表已存在\n")
        # 检测是否有用户存在
        for index in c.execute("SELECT * FROM user"):
            if index[0][0] != "":
                config.set('user', 'exist', 'True')
                config.write(open('config.ini', 'w'))
                break
            # 防止从其他地方拷贝ini文件，而数据库无用户导致程序崩溃
            else:
                config.set('user', 'exist', 'False')
                config.write(open('config.ini', 'w'))
                break
    user = User()
    widget = Widget()
    UI_U = user.ui
    user.init()
    if user.exec():
        appstart = datetime.datetime.now()
        # 抓图对象
        sc = imgGrab()
        UI = widget.ui
        widget.init()
        widget.updateTable()
        widget.show()
        try:
            # 加载每日一句
            loadsentences()
        except Exception as e:
            errorbox.write(str(datetime.datetime.now()) + '\n')
            errorbox.write("网络故障，请检查网络连通性\n")
        sys.exit(app.exec())
